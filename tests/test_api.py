import sys
from io import BytesIO
from pathlib import Path

from docx import Document
from fastapi.testclient import TestClient
from pypdf import PdfWriter


ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "src"))

from sci402_agent.api import create_app
from sci402_agent.feedback_agent import (
    MODE_1_SUPPORTIVE_INQUIRY,
    MODE_2_STRUCTURED_GUIDANCE,
)
from sci402_agent.llm_client import LLMCallError, LLMConfigurationError
from sci402_agent.llm_rubric_scorer import (
    LLMScoringValidationError,
    fallback_scoring_response,
)


def test_health_endpoint_reports_service_status(monkeypatch):
    client = TestClient(create_app())

    response = client.get("/health")

    assert response.status_code == 200
    assert response.json() == {"status": "ok", "criteria_count": 5}


def test_index_page_is_served(monkeypatch):
    client = TestClient(create_app())

    response = client.get("/")

    assert response.status_code == 200
    assert "text/html" in response.headers["content-type"]
    assert "SCI402 Feedback Agent" in response.text
    assert "Feedback style" in response.text
    assert "Score source" in response.text
    assert 'id="feedbackStyleValue">None</strong>' in response.text
    assert "Not selected" not in response.text
    assert "<strong>Local</strong>" in response.text
    assert "<strong>AI</strong>" in response.text
    assert "<strong>Fallback</strong>" in response.text
    assert "<strong>Guidance</strong>" in response.text
    assert "Structured guidance" not in response.text
    assert "AI semantic judgement" not in response.text
    assert 'id="helpButton"' in response.text
    assert 'id="helpModal"' in response.text
    assert 'id="uploadButton"' in response.text
    assert 'id="draftFileInput"' in response.text
    assert 'accept=".txt,.docx,.pdf"' in response.text
    assert "Analysis Modes" in response.text


def test_static_app_script_is_served(monkeypatch):
    client = TestClient(create_app())

    response = client.get("/static/app.js")

    assert response.status_code == 200
    assert "runFeedback" in response.text
    assert "runLLMScore" in response.text
    assert "runUpload" in response.text
    assert "postFormData" in response.text
    assert "FormData" in response.text
    assert '"/upload-draft"' in response.text
    assert "uploadButton" in response.text
    assert "draftFileInput" in response.text
    assert "resetAnalysisDisplay" in response.text
    assert "lastAIReview" in response.text
    assert "includeAIReview" in response.text
    assert "openHelpModal" in response.text
    assert "closeHelpModal" in response.text
    assert 'MODE_2_STRUCTURED_GUIDANCE: "Guidance"' in response.text
    assert "Structured guidance" not in response.text
    assert "AI semantic judgement" in response.text
    assert "AI second review" in response.text
    assert "scientific_reasoning_concerns" in response.text
    assert "local_precheck_blind_spots" in response.text
    assert "why_score_differs_from_local" in response.text
    assert "score-delta" in response.text
    assert "toggleCriterionCard" in response.text
    assert 'aria-expanded="false"' in response.text
    assert "criterion-card-body" in response.text
    assert "analysis.suggested_feedback_mode" in response.text
    assert "analysis.validated_scores || analysis.criterion_scores" in response.text


def test_static_styles_prioritize_score_source_width(monkeypatch):
    client = TestClient(create_app())

    response = client.get("/static/styles.css")

    assert response.status_code == 200
    assert "grid-template-columns: 1.02fr 1.28fr 0.58fr 0.74fr 0.62fr" in response.text
    assert "grid-column: span 2" not in response.text
    assert ".summary-strip > div:nth-child(n + 3)" in response.text
    assert "text-align: center" in response.text
    assert ".panel-actions" in response.text
    summary_start = response.text.index(".summary-strip strong")
    summary_end = response.text.index(".coverage-track")
    summary_css = response.text[summary_start:summary_end]
    assert "overflow-wrap: break-word" in summary_css
    assert "word-break: normal" in summary_css
    assert "overflow-wrap: anywhere" not in summary_css
    assert "white-space: nowrap" in summary_css


def test_static_styles_support_collapsible_criterion_cards(monkeypatch):
    client = TestClient(create_app())

    response = client.get("/static/styles.css")

    assert response.status_code == 200
    assert ".criterion-card-head" in response.text
    assert ".criterion-card-body" in response.text
    assert ".toggle-label" in response.text
    assert "overflow: hidden" in response.text


def test_criteria_endpoint_returns_ordered_rubric_rules(monkeypatch):
    client = TestClient(create_app())

    response = client.get("/criteria")

    assert response.status_code == 200
    payload = response.json()
    assert len(payload) == 5
    assert payload[0]["id"] == "C1_SCIENTIFIC_BACKGROUND"
    assert payload[0]["title"] == "Scientific Background & Problem Definition"


def test_single_criterion_endpoint_returns_404_for_unknown_id(monkeypatch):
    client = TestClient(create_app())

    response = client.get("/criteria/UNKNOWN")

    assert response.status_code == 404
    assert "Unknown criterion id" in response.json()["detail"]


def test_analyze_endpoint_returns_profile_and_criterion_coverage(monkeypatch):
    client = TestClient(create_app())

    response = client.post(
        "/analyze",
        json={
            "student_text": (
                "I will use a regression model with input features and an output target."
            )
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["word_count"] == 13
    assert payload["suggested_feedback_mode"] == MODE_1_SUPPORTIVE_INQUIRY
    assert payload["matched_criteria"] == ["C2_AI_ML_FORMULATION"]
    assert payload["criterion_coverage"][1]["id"] == "C2_AI_ML_FORMULATION"
    assert payload["criterion_coverage"][1]["is_matched"] is True
    assert payload["criterion_coverage"][0]["missing_feedback"]
    assert payload["estimated_total"] == payload["estimated_total_25"]
    assert payload["grade_band"]
    assert payload["structure_check"]["required_word_count"] == 3000
    assert len(payload["criterion_scores"]) == 5
    assert payload["priority_revisions"]


def test_upload_draft_txt_returns_extracted_text(monkeypatch):
    client = TestClient(create_app())

    response = client.post(
        "/upload-draft",
        files={"file": ("draft.txt", b"\xef\xbb\xbfProposal text here.", "text/plain")},
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["filename"] == "draft.txt"
    assert payload["extracted_text"] == "Proposal text here."
    assert payload["character_count"] == len("Proposal text here.")
    assert payload["warnings"] == []


def test_upload_draft_docx_returns_paragraph_and_table_text(monkeypatch):
    client = TestClient(create_app())
    document = Document()
    document.add_paragraph("Paragraph proposal text.")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Input features"
    table.cell(0, 1).text = "Output target"
    buffer = BytesIO()
    document.save(buffer)

    response = client.post(
        "/upload-draft",
        files={
            "file": (
                "draft.docx",
                buffer.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert "Paragraph proposal text." in payload["extracted_text"]
    assert "Input features | Output target" in payload["extracted_text"]
    assert payload["character_count"] == len(payload["extracted_text"])


def test_upload_draft_rejects_unsupported_file_type(monkeypatch):
    client = TestClient(create_app())

    response = client.post(
        "/upload-draft",
        files={"file": ("draft.rtf", b"{\\rtf1 draft}", "application/rtf")},
    )

    assert response.status_code == 400
    assert "Unsupported file type" in response.json()["detail"]


def test_upload_draft_empty_file_returns_warning(monkeypatch):
    client = TestClient(create_app())

    response = client.post(
        "/upload-draft",
        files={"file": ("empty.txt", b"", "text/plain")},
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["extracted_text"] == ""
    assert payload["character_count"] == 0
    assert "Uploaded file is empty." in payload["warnings"]


def test_upload_draft_blank_pdf_returns_scanned_warning(monkeypatch):
    client = TestClient(create_app())
    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    buffer = BytesIO()
    writer.write(buffer)

    response = client.post(
        "/upload-draft",
        files={"file": ("blank.pdf", buffer.getvalue(), "application/pdf")},
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["extracted_text"] == ""
    assert payload["character_count"] == 0
    assert any("No extractable text" in warning for warning in payload["warnings"])


def test_chat_endpoint_returns_model_content(monkeypatch):
    def fake_chat_completion(messages, tools=None):
        assert messages == [{"role": "user", "content": "Hello!"}]
        assert tools is None
        return {"content": "Hi there.", "tool_calls": []}

    monkeypatch.setattr("sci402_agent.api.chat_completion", fake_chat_completion)
    client = TestClient(create_app())

    response = client.post(
        "/chat",
        json={"messages": [{"role": "user", "content": "Hello!"}]},
    )

    assert response.status_code == 200
    assert response.json() == {"content": "Hi there.", "tool_calls": []}


def test_feedback_endpoint_returns_mode_analysis_and_feedback(monkeypatch):
    captured_request = {}

    def fake_chat_completion(messages, tools=None):
        captured_request["messages"] = messages
        assert tools is None
        return {
            "content": "Mode: MODE_2_STRUCTURED_GUIDANCE\nFeedback: Add workflow details.",
            "tool_calls": [],
        }

    monkeypatch.setattr("sci402_agent.api.chat_completion", fake_chat_completion)
    client = TestClient(create_app())

    response = client.post(
        "/feedback",
        json={
            "student_text": (
                "My project will use a regression model with input features "
                "and an output target for a dataset, but the workflow and "
                "validation are not planned yet."
            )
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["mode"] == MODE_2_STRUCTURED_GUIDANCE
    assert payload["analysis"]["matched_criteria"] == [
        "C2_AI_ML_FORMULATION",
        "C3_METHODOLOGY_WORKFLOW",
    ]
    assert payload["feedback"] == (
        "Mode: MODE_2_STRUCTURED_GUIDANCE\nFeedback: Add workflow details."
    )
    assert captured_request["messages"][0]["role"] == "system"
    assert "SCI402 Rubric Rules" in captured_request["messages"][0]["content"]
    assert "Priority rewrite targets" in captured_request["messages"][0]["content"]
    assert "Rubric diagnosis" not in captured_request["messages"][0]["content"]


def test_feedback_endpoint_accepts_ai_second_review(monkeypatch):
    captured_request = {}

    def fake_chat_completion(messages, tools=None):
        captured_request["messages"] = messages
        assert tools is None
        return {"content": "AI review-based feedback.", "tool_calls": []}

    monkeypatch.setattr("sci402_agent.api.chat_completion", fake_chat_completion)
    client = TestClient(create_app())

    ai_review = [
        {
            "id": "C2_AI_ML_FORMULATION",
            "title": "AI / ML Problem Formulation",
            "score_0_to_5": 4,
            "level": "Satisfactory",
            "local_score_0_to_5": 3,
            "evidence": ["input features and an output target"],
            "missing_items": ["algorithm justification"],
            "rationale": "The AI task is mostly clear.",
            "semantic_diagnosis": "The task framing is coherent but thin.",
            "quality_concerns": ["The model choice is not justified."],
            "scientific_reasoning_concerns": [
                "The model choice is not linked to the scientific data pattern."
            ],
            "local_precheck_blind_spots": [
                "Keyword coverage does not prove the output is measurable."
            ],
            "why_score_differs_from_local": (
                "AI gives extra credit for coherent task framing."
            ),
            "revision_focus": "Justify the algorithm against the data structure.",
            "confidence": "medium",
            "cap_applied": False,
            "blocking_flags": [],
            "invalid_evidence": [],
            "adjustments": [],
            "source": "llm",
        }
    ]

    response = client.post(
        "/feedback",
        json={
            "student_text": (
                "This regression project has input features and an output target."
            ),
            "ai_review": ai_review,
        },
    )

    assert response.status_code == 200
    assert response.json()["feedback"] == "AI review-based feedback."
    user_prompt = captured_request["messages"][1]["content"]
    assert "AI second review:" in user_prompt
    assert "scientific_reasoning_concerns" in user_prompt
    assert "local_precheck_blind_spots" in user_prompt
    assert "why_score_differs_from_local" in user_prompt
    assert "candidate_evidence_and_guardrails" not in user_prompt


def test_feedback_endpoint_returns_503_for_missing_llm_config(monkeypatch):
    def fake_chat_completion(messages, tools=None):
        raise LLMConfigurationError("Missing SCI402_LLM_API_KEY.")

    monkeypatch.setattr("sci402_agent.api.chat_completion", fake_chat_completion)
    client = TestClient(create_app())

    response = client.post(
        "/feedback",
        json={"student_text": "This is a proposal draft with enough words to analyze."},
    )

    assert response.status_code == 503
    assert response.json()["detail"] == "Missing SCI402_LLM_API_KEY."


def test_feedback_endpoint_returns_502_for_llm_call_error(monkeypatch):
    def fake_chat_completion(messages, tools=None):
        raise LLMCallError("Model call failed.")

    monkeypatch.setattr("sci402_agent.api.chat_completion", fake_chat_completion)
    client = TestClient(create_app())

    response = client.post(
        "/feedback",
        json={"student_text": "This is a proposal draft with enough words to analyze."},
    )

    assert response.status_code == 502
    assert response.json()["detail"] == "Model call failed."


def test_llm_score_endpoint_returns_validated_scores(monkeypatch):
    def fake_score_with_llm(student_text, profile):
        payload = fallback_scoring_response(profile, "fake fallback")
        payload["source"] = "llm"
        payload["fallback_reason"] = None
        payload["llm_scores"] = [
            {
                "id": score["id"],
                "score_0_to_5": score["score_0_to_5"],
                "evidence": score["evidence"],
                "missing_items": score["missing_items"],
                "rationale": "Mock LLM rationale.",
                "semantic_diagnosis": "Mock semantic diagnosis.",
                "quality_concerns": ["Mock quality concern."],
                "scientific_reasoning_concerns": ["Mock reasoning concern."],
                "local_precheck_blind_spots": ["Mock blind spot."],
                "why_score_differs_from_local": "Mock score difference.",
                "revision_focus": "Mock revision focus.",
                "confidence": "medium",
                "cap_applied": False,
            }
            for score in profile["criterion_scores"]
        ]
        return payload

    monkeypatch.setattr("sci402_agent.api.score_with_llm", fake_score_with_llm)
    client = TestClient(create_app())

    response = client.post(
        "/llm-score",
        json={
            "student_text": (
                "I will use a regression model with input features and an output target."
            )
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["source"] == "llm"
    assert payload["fallback_reason"] is None
    assert payload["local_analysis"]["word_count"] == 13
    assert len(payload["llm_scores"]) == 5
    assert len(payload["validated_scores"]) == 5
    assert payload["llm_scores"][0]["semantic_diagnosis"]
    assert payload["llm_scores"][0]["scientific_reasoning_concerns"]
    assert payload["validated_scores"][0]["local_score_0_to_5"] >= 0
    assert payload["validated_scores"][0]["local_precheck_blind_spots"]
    assert payload["validated_scores"][0]["why_score_differs_from_local"]
    assert payload["validated_scores"][0]["revision_focus"]
    assert payload["final_total"] == payload["final_total_25"]


def test_llm_score_endpoint_falls_back_when_llm_scoring_is_invalid(monkeypatch):
    def fake_score_with_llm(student_text, profile):
        raise LLMScoringValidationError("LLM returned invalid JSON.")

    monkeypatch.setattr("sci402_agent.api.score_with_llm", fake_score_with_llm)
    client = TestClient(create_app())

    response = client.post(
        "/llm-score",
        json={
            "student_text": (
                "I will use a regression model with input features and an output target."
            )
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["source"] == "local_fallback"
    assert payload["fallback_reason"] == "LLM returned invalid JSON."
    assert payload["llm_scores"] == []
    assert len(payload["validated_scores"]) == 5
