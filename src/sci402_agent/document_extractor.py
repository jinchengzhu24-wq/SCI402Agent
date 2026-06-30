"""In-memory draft text extraction for uploaded student files."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from pypdf import PdfReader


SUPPORTED_DRAFT_EXTENSIONS = {".txt", ".docx", ".pdf"}


class DraftExtractionError(ValueError):
    """Raised when an uploaded draft cannot be extracted."""


class UnsupportedDraftFileError(DraftExtractionError):
    """Raised when an uploaded draft extension is not supported."""


def extract_draft_text(filename: str, content: bytes) -> dict[str, Any]:
    """Extract text from a supported draft file without writing it to disk."""
    suffix = Path(filename or "").suffix.lower()
    warnings: list[str] = []

    if suffix not in SUPPORTED_DRAFT_EXTENSIONS:
        raise UnsupportedDraftFileError(
            "Unsupported file type. Upload a .txt, .docx, or .pdf draft."
        )

    if not content:
        return {
            "filename": filename,
            "character_count": 0,
            "extracted_text": "",
            "warnings": ["Uploaded file is empty."],
        }

    if suffix == ".txt":
        extracted_text = _extract_txt(content, warnings)
    elif suffix == ".docx":
        extracted_text = _extract_docx(content)
    else:
        extracted_text = _extract_pdf(content, warnings)

    extracted_text = extracted_text.strip()
    if not extracted_text:
        warnings.append(
            "No extractable text was found. If this is a scanned PDF, OCR is not supported yet."
        )

    return {
        "filename": filename,
        "character_count": len(extracted_text),
        "extracted_text": extracted_text,
        "warnings": warnings,
    }


def _extract_txt(content: bytes, warnings: list[str]) -> str:
    try:
        return content.decode("utf-8-sig")
    except UnicodeDecodeError:
        warnings.append(
            "Some characters could not be decoded as UTF-8 and were replaced."
        )
        return content.decode("utf-8", errors="replace")


def _extract_docx(content: bytes) -> str:
    try:
        document = Document(BytesIO(content))
    except Exception as exc:  # pragma: no cover - library error type is variable
        raise DraftExtractionError("Could not read the .docx file.") from exc

    chunks: list[str] = []
    chunks.extend(paragraph.text.strip() for paragraph in document.paragraphs)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                chunks.append(" | ".join(cells))

    return "\n".join(chunk for chunk in chunks if chunk)


def _extract_pdf(content: bytes, warnings: list[str]) -> str:
    try:
        reader = PdfReader(BytesIO(content))
    except Exception as exc:  # pragma: no cover - library error type is variable
        raise DraftExtractionError("Could not read the .pdf file.") from exc

    page_texts = []
    for page_number, page in enumerate(reader.pages, start=1):
        try:
            page_text = page.extract_text() or ""
        except Exception:
            warnings.append(f"Could not extract text from PDF page {page_number}.")
            page_text = ""
        if page_text.strip():
            page_texts.append(page_text.strip())

    return "\n\n".join(page_texts)
